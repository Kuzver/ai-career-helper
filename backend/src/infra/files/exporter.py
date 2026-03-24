from io import BytesIO
from html.parser import HTMLParser
from loguru import logger


def export_markdown(text: str) -> bytes:
    return text.encode("utf-8")


def export_html(text: str) -> bytes:
    import markdown
    html_body = markdown.markdown(text, extensions=["tables", "fenced_code"])

    html = f"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>ИИ-ассистент по карьере</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
  *,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
  :root{{
    --bg:#ffffff;--bg2:#f8f9ff;--bg3:#f1f3f9;
    --border:#e5e7eb;--text:#111827;--text2:#4b5563;--text3:#9ca3af;
    --brand:#3649F9;--brand-dim:rgba(54,73,249,0.08);--brand-border:rgba(54,73,249,0.25);
  }}
  .dark{{
    --bg:#0f172a;--bg2:#1e293b;--bg3:#334155;
    --border:#334155;--text:#f1f5f9;--text2:#94a3b8;--text3:#64748b;
    --brand:#818cf8;--brand-dim:rgba(129,140,248,0.12);--brand-border:rgba(129,140,248,0.3);
  }}
  body{{font-family:'Inter',system-ui,sans-serif;background:var(--bg);color:var(--text);min-height:100vh;padding:40px 20px 80px;transition:background .3s,color .3s}}
  .container{{max-width:800px;margin:0 auto}}
  .header{{display:flex;align-items:center;justify-content:space-between;margin-bottom:32px;padding-bottom:16px;border-bottom:1px solid var(--border)}}
  .logo{{display:flex;align-items:center;gap:10px}}
  .logo-circle{{width:32px;height:32px;border-radius:50%;background:#3649F9}}
  .logo-text{{font-size:13px;font-weight:600;color:var(--text)}}
  .theme-btn{{width:32px;height:32px;border-radius:8px;border:1px solid var(--border);background:var(--bg2);cursor:pointer;display:flex;align-items:center;justify-content:center;color:var(--text2)}}
  .theme-btn:hover{{border-color:var(--brand);color:var(--brand)}}
  .content{{line-height:1.8}}
  .content h1{{font-size:24px;font-weight:700;color:var(--text);margin:32px 0 12px}}
  .content h2{{font-size:18px;font-weight:600;color:var(--text);margin:28px 0 10px;padding-left:14px;border-left:3px solid var(--brand)}}
  .content h3{{font-size:15px;font-weight:600;color:var(--text);margin:20px 0 8px}}
  .content p{{font-size:14px;color:var(--text2);margin:8px 0}}
  .content strong{{color:var(--text);font-weight:600}}
  .content ul,.content ol{{padding-left:24px;margin:8px 0}}
  .content li{{font-size:14px;color:var(--text2);margin:4px 0}}
  .content code{{background:var(--brand-dim);color:var(--brand);padding:2px 6px;border-radius:4px;font-size:13px}}
  .content pre{{background:var(--bg3);border:1px solid var(--border);border-radius:10px;padding:16px;overflow-x:auto;margin:12px 0}}
  .content pre code{{background:none;color:var(--text2);padding:0}}
  .content blockquote{{background:var(--bg2);border-left:3px solid var(--brand);border-radius:0 10px 10px 0;padding:12px 16px;margin:12px 0;color:var(--text2)}}
  .content table{{width:100%;border-collapse:collapse;margin:12px 0;font-size:13px}}
  .content th{{background:var(--brand);color:#fff;font-weight:600;padding:8px 12px;text-align:left}}
  .content th:first-child{{border-radius:8px 0 0 0}}
  .content th:last-child{{border-radius:0 8px 0 0}}
  .content td{{padding:8px 12px;border-bottom:1px solid var(--border);color:var(--text2)}}
  .content tr:nth-child(even) td{{background:var(--bg2)}}
  .content a{{color:var(--brand);text-decoration:underline}}
  .content hr{{border:none;border-top:1px solid var(--border);margin:24px 0}}
  .footer{{margin-top:48px;padding-top:16px;border-top:1px solid var(--border);font-size:11px;color:var(--text3);display:flex;justify-content:space-between}}
  @media(max-width:540px){{body{{padding:20px 12px 40px}}}}
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <div class="logo">
      <div class="logo-circle"></div>
      <span class="logo-text">ИИ-ассистент по карьере</span>
    </div>
    <button class="theme-btn" onclick="document.documentElement.classList.toggle('dark')" title="Тема">&#9788;</button>
  </div>
  <div class="content">
{html_body}
  </div>
  <div class="footer">
    <span>Сгенерировано ИИ-ассистентом по карьере</span>
    <span>{__import__('datetime').date.today().strftime('%d.%m.%Y')}</span>
  </div>
</div>
</body>
</html>"""
    return html.encode("utf-8")


class _DocxBuilder(HTMLParser):
    """Парсит HTML и строит DOCX через python-docx."""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self._tag_stack: list[str] = []
        self._current_paragraph = None
        self._in_pre = False
        self._in_strong = False
        self._in_em = False
        self._in_blockquote = False
        self._list_stack: list[str] = []
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] | None = None
        self._cell_text = ""
        self._in_table = False
        self._in_th = False
        self._code_text = ""

    def handle_starttag(self, tag: str, attrs):
        self._tag_stack.append(tag)

        if tag in ("h1", "h2", "h3", "h4"):
            level = int(tag[1])
            self._current_paragraph = self.doc.add_heading("", level=level)

        elif tag == "p":
            if self._in_blockquote:
                self._current_paragraph = self.doc.add_paragraph()
            else:
                self._current_paragraph = self.doc.add_paragraph()

        elif tag == "ul":
            self._list_stack.append("ul")
        elif tag == "ol":
            self._list_stack.append("ol")

        elif tag == "li":
            if self._list_stack and self._list_stack[-1] == "ol":
                self._current_paragraph = self.doc.add_paragraph(style="List Number")
            else:
                self._current_paragraph = self.doc.add_paragraph(style="List Bullet")

        elif tag == "pre":
            self._in_pre = True
            self._code_text = ""

        elif tag == "code":
            pass

        elif tag in ("strong", "b"):
            self._in_strong = True
        elif tag in ("em", "i"):
            self._in_em = True

        elif tag == "blockquote":
            self._in_blockquote = True

        elif tag == "table":
            self._in_table = True
            self._table_rows = []

        elif tag == "tr":
            self._current_row = []

        elif tag in ("td", "th"):
            self._in_th = tag == "th"
            self._cell_text = ""

    def handle_endtag(self, tag: str):
        if self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()

        if tag in ("h1", "h2", "h3", "h4", "p", "li"):
            self._current_paragraph = None

        elif tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()

        elif tag == "pre":
            self._in_pre = False
            p = self.doc.add_paragraph()
            from docx.shared import Pt, RGBColor
            run = p.add_run(self._code_text.rstrip("\n"))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            self._code_text = ""

        elif tag == "code":
            pass

        elif tag in ("strong", "b"):
            self._in_strong = False
        elif tag in ("em", "i"):
            self._in_em = False

        elif tag == "blockquote":
            self._in_blockquote = False

        elif tag in ("td", "th"):
            if self._current_row is not None:
                self._current_row.append(self._cell_text)
            self._cell_text = ""

        elif tag == "tr":
            if self._current_row is not None:
                self._table_rows.append(self._current_row)
            self._current_row = None

        elif tag == "table":
            self._flush_table()
            self._in_table = False

    def handle_data(self, data: str):
        if self._in_pre:
            self._code_text += data
            return

        if self._in_table and (self._current_row is not None):
            self._cell_text += data
            return

        if self._current_paragraph is not None:
            from docx.shared import Pt
            tag_in_code = "code" in self._tag_stack
            run = self._current_paragraph.add_run(data)
            if tag_in_code:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            if self._in_strong:
                run.bold = True
            if self._in_em:
                run.italic = True

    def _flush_table(self):
        if not self._table_rows:
            return
        cols = max(len(r) for r in self._table_rows)
        rows_count = len(self._table_rows)
        table = self.doc.add_table(rows=rows_count, cols=cols)
        table.style = "Table Grid"
        from docx.shared import Pt
        for i, row_data in enumerate(self._table_rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    cell = row.cells[j]
                    cell.text = cell_text.strip()
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        self._table_rows = []


def md_to_docx(text: str, title: str = "") -> bytes:
    """Конвертирует Markdown в DOCX по ГОСТ (Times New Roman 14, 1.5 интервал, отступы)."""
    import markdown
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ГОСТ: поля 3cm left, 1.5cm right, 2cm top/bottom
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ГОСТ: Times New Roman 14pt, 1.5 межстрочный
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Cm(1.25)

    # Заголовки
    for i in range(1, 5):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Times New Roman"
        hs.font.size = Pt(16 - i)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)
        hs.paragraph_format.first_line_indent = Cm(0)

    if title:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tp.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"
        doc.add_paragraph()

    html_body = markdown.markdown(text, extensions=["tables", "fenced_code"])
    builder = _DocxBuilder(doc)
    builder.feed(html_body)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_docx(text: str) -> bytes:
    try:
        return md_to_docx(text)
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        return export_markdown(text)
